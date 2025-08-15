from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from typing import Dict, Any, List

def _set_normal_style(doc: Document, font_name="Times New Roman", size_pt=12):
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(size_pt)

def _heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(16)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        run.font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def _para(doc: Document, text: str, bold=False, italic=False, align_center=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align_center else WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def _numbered_list(doc: Document, items: List[str]):
    for i, it in enumerate(items, 1):
        _para(doc, f"{i}. {it}")

def _add_page_number_footer(doc: Document):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Page "); _add_field(run, "PAGE")
    run = p.add_run(" of "); _add_field(run, "NUMPAGES")

def _add_field(run, field_code: str):
    r = run._r
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve'); instrText.text = f" {field_code} "
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar'); fldChar3.set(qn('w:fldCharType'), 'end')
    r.append(fldChar1); r.append(instrText); r.append(fldChar2); r.append(fldChar3)

def build_docx_from_draft(draft: Dict[str, Any], out_path: str, title: str = "Draft"):
    doc = Document()
    _set_normal_style(doc); _add_page_number_footer(doc)
    _heading(doc, draft.get("title") or title, level=1)
    for label in [("Parties","parties"),("Facts","facts"),("Grounds","grounds"),("Prayer","prayer")]:
        items = draft.get(label[1]) or []
        if items:
            _heading(doc, label[0], level=2)
            _numbered_list(doc, items)
    for label in [("Annexures","annexures"),("Citations","citations")]:
        items = draft.get(label[1]) or []
        if items:
            _heading(doc, label[0], level=2)
            for it in items: _para(doc, f"- {it}")
    notes = draft.get("notes") or ""
    if notes:
        _heading(doc, "Notes", level=2); _para(doc, notes)
    doc.save(out_path)
